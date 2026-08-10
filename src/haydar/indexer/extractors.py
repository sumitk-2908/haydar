"""
Text extraction module for Haydar indexer.
"""

import logging
from dataclasses import dataclass, field
from enum import Enum
from pathlib import Path
from typing import Any

import chardet
import docx
import pypdf

from haydar.config import (
    CACHE_DIR,
    CODE_EXTENSIONS,
    IMAGE_EXTENSIONS,
    TEXT_EXTENSIONS,
)
from haydar.ocr import TesseractStatus, detect_tesseract

logger = logging.getLogger(__name__)


class Disposition(Enum):
    """What happened to a file, recorded rather than inferred.

    The distinction that matters most is ``OCR_DEFERRED`` vs ``EMPTY``: an image
    Haydar could not read because OCR is unavailable is *not* an image that
    contained no text. Only the latter is finished work.
    """

    CONTENT = "content"
    EMPTY = "empty"
    UNCHANGED = "unchanged"
    UNSUPPORTED = "unsupported"
    TOO_LARGE = "too_large"
    OCR_DEFERRED = "ocr_deferred"
    TRANSIENT_ERROR = "transient_error"
    PERMANENT_ERROR = "permanent_error"


@dataclass
class ExtractedContent:
    """Dataclass holding the extracted text and associated metadata."""
    text: str
    metadata: dict[str, Any] = field(default_factory=dict)


@dataclass(frozen=True)
class ExtractionOutcome:
    """A typed extraction result: never a bare ``None`` the caller must guess at."""

    disposition: Disposition
    content: ExtractedContent | None = None
    detail: str = ""
    # Set for images so a later OCR upgrade can schedule an image-only refresh.
    ocr_engine_version: str = ""

    @property
    def text(self) -> str:
        return self.content.text if self.content is not None else ""

    @property
    def is_complete(self) -> bool:
        """Whether this outcome may be cached as finished work for this version."""
        return self.disposition in (
            Disposition.CONTENT,
            Disposition.EMPTY,
            Disposition.UNSUPPORTED,
        )


# Upper bound for the on-disk extraction cache (CACHE_DIR/*.txt). When the total
# size exceeds this, the oldest entries (by mtime) are evicted until it fits.
EXTRACTION_CACHE_MAX_BYTES = 512 * 1024 * 1024  # 512 MB


def prune_extraction_cache(max_bytes: int = EXTRACTION_CACHE_MAX_BYTES) -> int:
    """Evict oldest cached extraction files until the cache fits ``max_bytes``.

    Returns the number of bytes freed. Safe to call when the cache dir does not
    exist. Intended to run at ``init``/``reindex`` boundaries, not per-file.
    """
    if not CACHE_DIR.is_dir():
        return 0

    try:
        entries = []
        total = 0
        for path in CACHE_DIR.glob("*.txt"):
            try:
                stat = path.stat()
            except OSError:
                continue
            entries.append((stat.st_mtime, stat.st_size, path))
            total += stat.st_size

        if total <= max_bytes:
            return 0

        freed = 0
        # Oldest first.
        for _mtime, size, path in sorted(entries, key=lambda e: e[0]):
            if total - freed <= max_bytes:
                break
            try:
                path.unlink()
                freed += size
            except OSError:
                continue

        if freed:
            logger.info("Pruned extraction cache: freed %d bytes", freed)
        return freed
    except Exception as e:
        logger.warning("Failed to prune extraction cache: %s", e)
        return 0

def _extract_pdf(filepath: Path) -> ExtractedContent | None:
    try:
        reader = pypdf.PdfReader(filepath)
        if reader.is_encrypted:
            try:
                reader.decrypt("")
            except Exception:
                logger.warning(f"Failed to decrypt PDF (empty password failed): {filepath}")
                return None

        text_parts = []
        for page in reader.pages:
            try:
                text_parts.append(page.extract_text(extraction_mode="layout") or "")
            except Exception as e:
                logger.warning(f"Failed to extract text from a page in {filepath}: {e}")

        return ExtractedContent(
            text="\n".join(text_parts),
            metadata={"page_count": len(reader.pages)}
        )
    except Exception as e:
        logger.warning(f"Failed to extract PDF {filepath}: {e}")
        return None


def _extract_docx(filepath: Path) -> ExtractedContent | None:
    try:
        doc = docx.Document(str(filepath))
        text_parts = []

        # Headers
        for section in doc.sections:
            for header_para in section.header.paragraphs:
                text_parts.append(header_para.text)

        # Body paragraphs
        for para in doc.paragraphs:
            text_parts.append(para.text)

        # Tables
        for table in doc.tables:
            if hasattr(table, 'iter_inner_content'):
                for content in table.iter_inner_content():
                    text_parts.append(getattr(content, 'text', ''))
            else:
                for row in table.rows:
                    for cell in row.cells:
                        text_parts.append(cell.text)

        # Footers
        for section in doc.sections:
            for footer_para in section.footer.paragraphs:
                text_parts.append(footer_para.text)

        return ExtractedContent(
            text="\n".join(t for t in text_parts if t.strip()),
            metadata={}
        )
    except Exception as e:
        logger.warning(f"Failed to extract DOCX {filepath}: {e}")
        return None


def _extract_text(filepath: Path) -> ExtractedContent | None:
    try:
        raw = filepath.read_bytes()
        sample = raw[:10240]
        result = chardet.detect(sample)
        encoding = result['encoding']
        confidence = result['confidence']

        if not encoding or confidence < 0.5:
            encoding = 'utf-8'

        text = raw.decode(encoding, errors='replace')
        return ExtractedContent(
            text=text,
            metadata={"encoding": encoding}
        )
    except Exception as e:
        logger.warning(f"Failed to extract text {filepath}: {e}")
        return None


def _extract_image(filepath: Path) -> ExtractionOutcome:
    """OCR an image, distinguishing "no OCR available" from "no text found"."""
    info = detect_tesseract()
    if info.status is not TesseractStatus.FOUND:
        logger.debug("Image OCR unavailable (%s): %s", info.status.value, filepath)
        # Deferral, not completion. This file stays eligible so a later one-click
        # OCR install can index it without a full reindex.
        return ExtractionOutcome(
            disposition=Disposition.OCR_DEFERRED,
            detail=info.status.value,
        )

    engine_version = f"tesseract-{info.version}" if info.version else "tesseract"
    try:
        import pytesseract

        if info.path:
            pytesseract.pytesseract.tesseract_cmd = info.path
        text = pytesseract.image_to_string(str(filepath), lang="eng")
    except pytesseract.TesseractNotFoundError:
        logger.warning("Tesseract OCR engine became unavailable during extraction.")
        return ExtractionOutcome(
            disposition=Disposition.OCR_DEFERRED, detail="engine_disappeared"
        )
    except Exception as exc:
        logger.exception("Failed to extract image %s", filepath)
        return ExtractionOutcome(
            disposition=Disposition.TRANSIENT_ERROR,
            detail=str(exc)[:200],
            ocr_engine_version=engine_version,
        )

    if not text.strip():
        # OCR ran and genuinely found nothing. That is a finished result, and it
        # is tagged with the engine version so an upgrade can revisit it.
        return ExtractionOutcome(
            disposition=Disposition.EMPTY,
            content=ExtractedContent(text="", metadata={}),
            ocr_engine_version=engine_version,
        )
    return ExtractionOutcome(
        disposition=Disposition.CONTENT,
        content=ExtractedContent(text=text, metadata={"ocr": engine_version}),
        ocr_engine_version=engine_version,
    )


def extract(
    filepath: Path, file_hash: str | None = None, *, refresh_ocr: bool = False
) -> ExtractionOutcome:
    """Extract a file's text and report a typed disposition.

    This is the form the indexing engine uses: every outcome is explicit, so a
    caller can never confuse "OCR is not installed" with "this file is empty".

    ``refresh_ocr`` bypasses the extraction-text cache for images. The cache is
    keyed by file content, which is exactly wrong for a re-OCR: the bytes are
    unchanged, but the engine that reads them is new, so a cache hit would
    return the old engine's text and leave the image eligible forever.
    """
    ext = filepath.suffix.lower()
    is_image = ext in IMAGE_EXTENSIONS

    cache_path = None
    if file_hash:
        cache_path = CACHE_DIR / f"{file_hash}.txt"
        if cache_path.exists() and not (refresh_ocr and is_image):
            try:
                text = cache_path.read_text(encoding="utf-8")
                return ExtractionOutcome(
                    disposition=Disposition.CONTENT if text.strip() else Disposition.EMPTY,
                    content=ExtractedContent(text=text, metadata={"cached": True}),
                )
            except OSError:
                logger.debug("Ignoring unreadable extraction cache for %s", filepath)

    if is_image:
        outcome = _extract_image(filepath)
    elif ext == ".pdf":
        outcome = _wrap(_extract_pdf(filepath))
    elif ext == ".docx":
        outcome = _wrap(_extract_docx(filepath))
    elif ext in TEXT_EXTENSIONS or ext in CODE_EXTENSIONS:
        outcome = _wrap(_extract_text(filepath))
    else:
        logger.debug("Unsupported extension %s for %s", ext, filepath)
        return ExtractionOutcome(disposition=Disposition.UNSUPPORTED)

    if outcome.disposition is Disposition.CONTENT and cache_path is not None:
        try:
            cache_path.write_text(outcome.text, encoding="utf-8")
        except OSError as e:
            # A cache miss next run is cheap; failing the file would not be.
            logger.warning("Failed to write extraction cache for %s: %s", filepath, e)

    return outcome


def _wrap(result: ExtractedContent | None) -> ExtractionOutcome:
    """Adapt a legacy ``None``-on-failure extractor to a typed outcome."""
    if result is None:
        return ExtractionOutcome(disposition=Disposition.PERMANENT_ERROR)
    if not result.text.strip():
        return ExtractionOutcome(disposition=Disposition.EMPTY, content=result)
    return ExtractionOutcome(disposition=Disposition.CONTENT, content=result)


def extract_text(filepath: Path, file_hash: str | None = None) -> ExtractedContent | None:
    """Dispatch text extraction based on file extension.

    Compatibility wrapper around :func:`extract` for callers that only need the
    content. Prefer :func:`extract` in new code: this signature cannot express
    the difference between a deferral and an empty file.
    """
    outcome = extract(filepath, file_hash=file_hash)
    return outcome.content
